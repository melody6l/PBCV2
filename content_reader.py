"""
文件内容提取模块 — 读取各类文件的文字内容，用于 L3 内容匹配。

支持格式：
- 纯文本：.txt/.csv/.log/.json/.xml/.md 等 → 直接读取
- Word: .docx → python-docx 解析 → Markdown
- Excel: .xlsx/.xls → openpyxl 解析 → 结构化摘要
- PDF: .pdf → pdfplumber 提取文字，失败则 PyMuPDF 转图片走 OCR
- 图片: .png/.jpg/.jpeg/.bmp/.tiff/.gif/.webp → 百度 OCR
"""

import hashlib
import os
import base64
import json
import sys
from datetime import datetime
from io import BytesIO
import threading

import requests
from PIL import Image

# ─── 进度状态（线程安全）───
_progress_lock = threading.Lock()
_current_progress = {
    "running": False,
    "total": 0,
    "current": 0,
    "current_file": "",
    "current_status": "",  # "缓存命中" / "正在OCR…" / "正在提取…" / "跳过"
    "text_count": 0,
    "ocr_count": 0,
    "skip_count": 0,
    "errors": [],
    "done": False,
    "logs": [],  # [{time, file, action, status, detail}]
}


def reset_progress():
    """重置进度状态。"""
    global _current_progress
    with _progress_lock:
        _current_progress = {
            "running": False, "total": 0, "current": 0, "current_file": "",
            "current_status": "", "text_count": 0, "ocr_count": 0,
            "skip_count": 0, "errors": [], "done": False, "logs": [],
        }


def _log(file_name, action, status="", detail=""):
    """添加一条日志（线程安全）。"""
    t = datetime.now().strftime("%H:%M:%S")
    entry = {"time": t, "file": file_name, "action": action, "status": status, "detail": detail}
    with _progress_lock:
        _current_progress.setdefault("logs", []).append(entry)



def update_progress(**kw):
    """更新进度字段（线程安全）。"""
    with _progress_lock:
        _current_progress.update(kw)


def get_progress():
    """获取当前进度快照（线程安全）。"""
    with _progress_lock:
        return dict(_current_progress)

# ─── 截断限制 ───
MAX_TEXT_LENGTH = 30000
MAX_DOCX_LENGTH = 40000
MAX_PDF_LENGTH = 50000
MAX_OCR_TEXT = 1000       # OCR 噪声多，送入 LLM 的截断更短
MAX_CONTENT_FOR_LLM = {
    "ocr": 1000,
    "text": 2000,
    "markdown": 2000,
    "structured": 0,       # 已是摘要，不截断
}

# ─── 可读文本文件扩展名 ───
TEXT_EXTENSIONS = {
    ".txt", ".csv", ".log", ".json", ".xml", ".md", ".html", ".htm",
    ".ini", ".cfg", ".conf", ".yml", ".yaml", ".sh", ".bat", ".py",
    ".js", ".css", ".sql", ".rst", ".ts", ".tsx", ".jsx",
}

# ─── 图片文件扩展名（走 OCR）───
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".gif", ".webp"}

# ─── 百度 OCR API ───
BAIDU_TOKEN_URL = "https://aip.baidubce.com/oauth/2.0/token"
BAIDU_OCR_URL = "https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic"


def _file_md5(file_path):
    """分块计算文件 MD5，用于缓存 key。"""
    hasher = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def _get_baidu_token(api_key, secret_key):
    """获取百度 OCR access_token。"""
    resp = requests.post(
        BAIDU_TOKEN_URL,
        data={"grant_type": "client_credentials", "client_id": api_key, "client_secret": secret_key},
        timeout=15,
    )
    resp.raise_for_status()
    data = resp.json()
    if "access_token" in data:
        return data["access_token"]
    raise RuntimeError(f"百度OCR认证失败: {data.get('error_description', data)}")


def _ocr_image_bytes(image_bytes, access_token):
    """对图片字节进行百度 OCR，返回识别文字。"""
    if len(image_bytes) > 4 * 1024 * 1024:
        return ""  # 单张图片超过 4MB 跳过 OCR
    encoded = base64.b64encode(image_bytes).decode("utf-8")
    resp = requests.post(
        BAIDU_OCR_URL,
        data={"image": encoded, "detect_direction": "true"},
        params={"access_token": access_token},
        timeout=30,
    )
    resp.raise_for_status()
    data = resp.json()
    if "words_result" in data:
        lines = [item["words"] for item in data["words_result"]]
        return "\n".join(lines)
    if data.get("error_msg"):
        raise RuntimeError(f"百度OCR识别失败: {data['error_msg']}")
    return ""


# ─── 各类型处理器 ───

def _read_text_file(path):
    """读取纯文本文件，尝试多种编码。"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as f:
                text = f.read()
            return {
                "content": text[:MAX_TEXT_LENGTH],
                "content_type": "text",
                "content_label": "文本文件",
                "truncated": len(text) > MAX_TEXT_LENGTH,
                "error": None,
            }
        except (UnicodeDecodeError, LookupError):
            continue
    return {
        "content": "",
        "content_type": "text",
        "content_label": "文本文件(解码失败)",
        "truncated": False,
        "error": "无法解码文件内容",
    }


def _read_docx(path):
    """读取 .docx，转换为 Markdown。"""
    try:
        from docx import Document
    except ImportError:
        return {"content": "", "content_type": "markdown", "content_label": "Word文档",
                "truncated": False, "error": "未安装 python-docx"}

    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            parts.append("")
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            level = style_name.replace("Heading", "").strip()
            try:
                level = int(level)
            except ValueError:
                level = 1
            parts.append(f"{'#' * min(level, 6)} {text}")
        elif "Title" in style_name:
            parts.append(f"# {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            parts.append("")  # 表格前后空行
            parts.extend(rows[:20])  # 最多 20 行
            if len(rows) > 20:
                parts.append(f"| ... 共 {len(rows)} 行 |")
            parts.append("")

    full = "\n".join(parts)
    return {
        "content": full[:MAX_DOCX_LENGTH],
        "content_type": "markdown",
        "content_label": "Word文档",
        "truncated": len(full) > MAX_DOCX_LENGTH,
        "error": None,
    }


def _read_xlsx_summary(path):
    """读取 .xlsx，输出结构化摘要：Sheet名 + 列头 + 前 5 行。"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"content": "", "content_type": "structured", "content_label": "Excel文件",
                "truncated": False, "error": "未安装 openpyxl"}

    try:
        wb = load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        return {"content": "", "content_type": "structured", "content_label": "Excel文件",
                "truncated": False, "error": f"无法打开Excel: {e}"}

    sheets_output = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = list(ws.iter_rows(max_row=7, values_only=True))  # 1 header + 5 data + 1 extra
        if not rows:
            continue
        headers = [str(c) if c is not None else "" for c in rows[0]]
        data_rows = rows[1:6]  # 最多 5 行数据
        data_strs = [" | ".join(str(c) if c is not None else "" for c in row) for row in data_rows]

        parts = [f"[Sheet: {name}]"]
        parts.append(f"列: {', '.join(h for h in headers if h)}")
        if data_strs:
            parts.append(f"数据行(前{len(data_strs)}行):")
            parts.extend(f"  {ds}" for ds in data_strs)
        sheets_output.append("\n".join(parts))

    wb.close()
    content = "\n---\n".join(sheets_output)
    return {
        "content": content,
        "content_type": "structured",
        "content_label": "Excel概要",
        "truncated": False,
        "error": None,
    }


def _read_pdf(path, ocr_config=None):
    """读取 PDF：先尝试 pdfplumber 提取文字，失败则 OCR。"""
    # Step 1: pdfplumber 提取文字
    try:
        import pdfplumber
    except ImportError:
        return {"content": "", "content_type": "text", "content_label": "PDF文档",
                "truncated": False, "error": "未安装 pdfplumber"}

    try:
        all_text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text)
        text = "\n".join(all_text).strip()
        if len(text) >= 50:
            return {
                "content": text[:MAX_PDF_LENGTH],
                "content_type": "text",
                "content_label": "PDF文档",
                "truncated": len(text) > MAX_PDF_LENGTH,
                "error": None,
            }
    except Exception as e:
        pass  # pdfplumber 失败，fallback 到 OCR

    # Step 2: OCR fallback — 每页单独 try，一张失败不卡死
    if not ocr_config or not ocr_config.get("api_key") or not ocr_config.get("secret_key"):
        return {
            "content": "", "content_type": "ocr", "content_label": "PDF(未配置OCR)",
            "truncated": False, "error": "PDF无文字层且未配置OCR",
        }

    all_text = []
    errors = []
    try:
        token = _get_baidu_token(ocr_config["api_key"], ocr_config["secret_key"])
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        for i, page in enumerate(doc):
            try:
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                page_text = _ocr_image_bytes(img_bytes, token)
                if page_text:
                    all_text.append(page_text)
            except Exception as e:
                msg = f"第{i+1}页OCR失败: {e}"
                errors.append(msg)
        doc.close()
    except Exception as e:
        return {"content": "", "content_type": "ocr", "content_label": "PDF(OCR失败)",
                "truncated": False, "error": f"PDF OCR初始化失败: {e}"}

    content = "\n".join(all_text)
    err_msg = "; ".join(errors) if errors else None
    return {
        "content": content,
        "content_type": "ocr",
        "content_label": "PDF(OCR)",
        "truncated": False,
        "error": err_msg,
    }


def _ocr_image_file(path, ocr_config=None):
    """对图片文件进行百度 OCR。"""
    if not ocr_config or not ocr_config.get("api_key") or not ocr_config.get("secret_key"):
        return {
            "content": "", "content_type": "ocr", "content_label": "图片(未配置OCR)",
            "truncated": False, "error": "未配置百度OCR凭证",
        }

    # 校验图片
    try:
        img = Image.open(path)
        img.verify()
    except Exception as e:
        return {"content": "", "content_type": "ocr", "content_label": "图片(无效)",
                "truncated": False, "error": f"无法打开图片: {e}"}

    try:
        token = _get_baidu_token(ocr_config["api_key"], ocr_config["secret_key"])
        with open(path, "rb") as f:
            img_bytes = f.read()
        text = _ocr_image_bytes(img_bytes, token)
        return {
            "content": text,
            "content_type": "ocr",
            "content_label": "图片(OCR)",
            "truncated": False,
            "error": None,
        }
    except Exception as e:
        return {"content": "", "content_type": "ocr", "content_label": "图片(OCR失败)",
                "truncated": False, "error": f"OCR失败: {e}"}


# ─── 公共接口 ───

# 返回结果类型: dict with keys: content, content_type, content_label, truncated, error
ContentResult = dict


def extract_content(file_path, ocr_config=None):
    """
    提取单个文件的内容。

    Args:
        file_path: 文件绝对路径
        ocr_config: {"api_key": str, "secret_key": str} 或 None

    Returns:
        ContentResult: {content, content_type, content_label, truncated, error}
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext in TEXT_EXTENSIONS:
        return _read_text_file(file_path)
    if ext == ".docx":
        return _read_docx(file_path)
    if ext in (".xlsx", ".xls"):
        return _read_xlsx_summary(file_path)
    if ext == ".pdf":
        return _read_pdf(file_path, ocr_config)
    if ext in IMAGE_EXTENSIONS:
        return _ocr_image_file(file_path, ocr_config)

    # 不支持的类型
    return {
        "content": "", "content_type": None, "content_label": "不支持的文件类型",
        "truncated": False, "error": f"不支持的文件格式: {ext or '未知'}",
    }


def get_content_for_llm(result):
    """
    从 ContentResult 中截取适合送入 LLM 的内容片段。

    根据不同 content_type 应用不同截断策略：OCR 截 1000 字符（噪声多），
    text/markdown 截 2000 字符，structured 不截断。
    """
    ct = result.get("content_type") or ""
    limit = MAX_CONTENT_FOR_LLM.get(ct, 2000)
    content = result.get("content") or ""
    return content[:limit] if limit > 0 else content


def extract_contents(file_paths, ocr_config=None, cache=None):
    """
    批量提取文件内容（带缓存）。

    Args:
        file_paths: 文件绝对路径列表
        ocr_config: 百度 OCR 凭证
        cache: {md5: ContentResult} 或 None（就地读写）

    Returns:
        (results, new_cache)
        results: {file_path: ContentResult}
        new_cache: 更新后的缓存 dict
    """
    if cache is None:
        cache = {}

    MAX_FILES = 30           # 单次最多处理 30 个文件
    MAX_IMAGE_BYTES = 8 * 1024 * 1024  # 超过 8MB 的图片跳过 OCR（太慢）
    OCR_SKIP_LARGE = True

    total = min(len(file_paths), MAX_FILES)
    results = {}
    text_count = 0
    ocr_count = 0
    skip_count = 0
    errors = []

    reset_progress()
    update_progress(running=True, total=total, current=0, done=False)

    for i, path in enumerate(file_paths[:MAX_FILES]):
        basename = os.path.basename(path)
        ext = os.path.splitext(path)[1].lower()
        idx = i + 1

        update_progress(current=idx, current_file=basename, current_status="正在读取…")

        try:
            file_size = os.path.getsize(path)
        except OSError:
            file_size = 0

        # 超大图片直接跳过
        if OCR_SKIP_LARGE and ext in IMAGE_EXTENSIONS and file_size > MAX_IMAGE_BYTES:
            _log(basename, "跳过", "文件过大", f"{file_size/1024/1024:.1f}MB")
            update_progress(current_status="跳过(文件过大)")
            results[path] = {"content": "", "content_type": "ocr", "content_label": "图片(文件过大)",
                             "truncated": False, "error": f"图片过大({file_size/1024/1024:.1f}MB)，跳过OCR"}
            skip_count += 1
            continue

        try:
            md5 = _file_md5(path)
        except (OSError, IOError):
            _log(basename, "读取失败", "无法读取")
            update_progress(current_status="读取失败")
            results[path] = {"content": "", "content_type": None, "content_label": "读取失败",
                             "truncated": False, "error": "无法读取文件（计算MD5失败）"}
            errors.append(f"{basename}: 读取失败")
            continue

        # 缓存命中
        if md5 in cache:
            ct = cache[md5].get("content_type", "")
            label = cache[md5].get("content_label", ct)
            _log(basename, "缓存命中", label)
            update_progress(current_status=f"缓存命中({ct})")
            results[path] = dict(cache[md5])
            if cache[md5].get("content"):
                if ct == "ocr":
                    ocr_count += 1
                else:
                    text_count += 1
            continue

        # 正在提取（OCR 可能很慢，专门显示）
        is_ocr = ext in IMAGE_EXTENSIONS or (ext == ".pdf" and ocr_config)
        if is_ocr:
            _log(basename, "正在OCR识别…")
            update_progress(current_status="正在OCR识别…")
        else:
            _log(basename, "正在提取内容…")
            update_progress(current_status="正在提取内容…")

        result = extract_content(path, ocr_config)
        ct = result.get("content_type", "")
        label = result.get("content_label", ct)
        if result.get("error"):
            _log(basename, "提取失败", label, result["error"])
            update_progress(current_status="提取失败")
            errors.append(f"{basename}: {result['error']}")
        elif result.get("content"):
            content_len = len(result["content"])
            if ct == "ocr":
                ocr_count += 1
                _log(basename, "OCR完成", f"{content_len}字")
            else:
                text_count += 1
                _log(basename, "读取完成", label, f"{content_len}字")
            update_progress(current_status="完成")
        else:
            _log(basename, "无内容", label)

        # 存缓存
        cache_entry = {
            "content": result["content"],
            "content_type": result["content_type"],
            "content_label": result["content_label"],
            "cached_at": datetime.now().isoformat(),
            "md5": md5,
        }
        cache[md5] = cache_entry
        results[path] = dict(cache_entry)

    update_progress(running=False, done=True,
                    current=total, current_file="",
                    text_count=text_count, ocr_count=ocr_count,
                    skip_count=skip_count, errors=errors)
    sys.stdout.flush()

    return results, cache
